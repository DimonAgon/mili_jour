import enum
import re

import django.db.models
import htmldocx #TODO: move down
from aiogram import F
from aiogram import types
from aiogram.filters import Command, CommandObject
from aiogram.methods import GetChatAdministrators
from aiogram.fsm.context import FSMContext
from aiogram.filters.state import State, StatesGroup

from aiogram_forms import FormsManager
from aiogram_forms.errors import ValidationError

from .dispatcher import dp,\
    commands_router, journal_router,\
    presence_poll_router,\
    registration_router,\
    journal_registration_subrouter,\
    bot
from ..models import *
from .forms.forms import *
from ..db_actions import *
from .filters import *
from .middleware import *
from ..infrastructure.enums import *
from ..infrastructure.ScheduleTiming import *
from .validators import *
from .checks import *
from static_text.misc import *
from static_text.utilities import *
from static_text import chat_messages
from static_text.chat_messages import *
from static_text import logging_messages
from static_text.logging_messages import *
from logging_native.utilis.frame_log_track.frame_log_track import log_track_frame
from .logger import handlers_logger
from misc.re_patterns import *
from misc.exceptions import *
from .misc import *

import asyncio

import datetime

from django.utils import timezone

import docx

import tempfile, os

import random

from key_generator import key_generator

from typing import Any

import operator

#TODO: unify code

logger = handlers_logger

untracked_log_data = {
    'bot'             ,
    'event_from_user' ,
    'event_chat'      ,
    'fsm_storage'     ,
    'state'           ,
    'raw_state'       ,
    'handler'         ,
    'event_update'    ,
    'event_router'    ,
    'command'         ,
    'forms'           ,
}

dp.message.outer_middleware(PassUserCredentials())
dp.poll_answer.outer_middleware(PassUserCredentials())
dp.message.outer_middleware(PassChatCredentials())
dp.poll_answer.outer_middleware(PassChatCredentials())
commands_router.message.middleware(ApplyArguments())
presence_poll_router.poll_answer.filter(PresencePollFilter())
journal_router.message.middleware(SuperuserSetJournal())
journal_registration_subrouter.message.middleware(SuperuserSetJournal())

#TODO: check set-journal states impact robotability

prefixes = {'🗡', '/'} #TODO: move to a separate file prefixes.py
#TODO: add pauses to messaging
#TODO: simplify register commands' names-- remove "register" word
@commands_router.message(Command(commands='start')) #TODO add middleware to show help for superuser
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def start_command(message: types.Message, *args, **kwargs):  # Self-presentation of the bot
    logger.info('')

    await message.reply(start_chat_info_message)


@commands_router.message(Command(commands='help', prefix=prefixes))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def help_command(message: types.Message, *args, **kwargs):
    logger.info('')

    #TODO: on_update_info
    await message.reply(HELPFUL_REPLY)


#TODO: split to two different functions for each mode-group
def poll_time_interval(mode, lesson: Lesson=None, last_lesson: Lesson=None): #TODO: set till_lesson instead of bot lesson and last_lesson
    now = timezone.localtime(timezone.now())
    if mode == PresenceMode.LIGHT_MODE:
        last_lesson_time = ScheduleTiming.lessons_intervals[last_lesson.ordinal]
        if last_lesson_time.upper < now.time(): #TODO: add checker
            raise LessonSkippedException

        deadline_time = last_lesson_time.upper
        deadline = now.replace(hour=deadline_time.hour, minute=deadline_time.minute, second=deadline_time.second)

    if not mode == PresenceMode.LIGHT_MODE: #TODO: fix too-early poll-time, should be recess before first lesson
        now_time = timezone.localtime(timezone.now()).time()

        lesson_time_interval = ScheduleTiming.lessons_intervals[lesson.ordinal]
        if lesson_time_interval.contains(now_time): start_time = (now + datetime.timedelta(seconds=1)).time()
        elif now_time < lesson_time_interval.lower:  start_time = lesson_time_interval.lower
        else:
            raise LessonSkippedException

        end_time = lesson_time_interval.upper
        deadline = now.replace(hour=end_time.hour, minute=end_time.minute, second=end_time.second)


        if mode == PresenceMode.HARDCORE_MODE:
            lower = start_time
            upper = end_time
            lower_today = now.replace(hour=lower.hour, minute=lower.minute, second=lower.second)
            upper_today = now.replace(hour=upper.hour, minute=upper.minute, second=lower.second)
            lower_today_timestamp = lower_today.timestamp()
            upper_today_timestamp = upper_today.timestamp()
            lower_today_timestamp_integer = int(lower_today_timestamp)
            upper_today_timestamp_integer = int(upper_today_timestamp)
            random_datetime_timestamp_integer = random.randint(lower_today_timestamp_integer,
                                                               upper_today_timestamp_integer)
            random_lesson_datetime = datetime.datetime.fromtimestamp(random_datetime_timestamp_integer)

            poll_time = now.replace(hour=random_lesson_datetime.hour, minute=random_lesson_datetime.minute, second=random_lesson_datetime.second)

        else: poll_time = now.replace(hour=start_time.hour, minute=start_time.minute, second=start_time.second)

        return P.openclosed(poll_time, deadline)

    return P.openclosed(now, deadline)


@commands_router.message(Command(commands=['presence', 'p'], prefix=prefixes),
                         F.chat.type.in_({'group', 'supergroup'}),
                         IsAdminFilter(),
                         AftercommandFullCheck(allow_no_argument=True,
                                      modes=PresenceMode,
                                      mode_checking=True,
                                      allow_no_mode= True,
                                      )
                         )
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)#processing first two lessons written conjointly raise mode validation error
#TODO: configure superusage
async def presence_command(message: types.Message, mode=default, flag=None, *args, **kwargs): #TODO: modify L mode to send 'arrived' and 'left' polls to calculate presence data
    chat_id = message.chat.id
    journal = await get_journal_async({'external_id':chat_id})
    now = timezone.localtime(timezone.now())
    today = now.date()
    current_schedule_attributes = {'date': today, 'journal': journal}
    try:
        current_schedule = await get_current_schedule_async(current_schedule_attributes)
        logger.info(
            current_schedule_by_attributes_existence_check_success_logging_info_message.format(
                current_schedule_attributes=current_schedule_attributes
            )
        )

    except CurrentSchedule.DoesNotExist:
        logger.error(
            current_schedule_by_attributes_existence_check_fail_logging_error_message.format(
                current_schedule_attributes=current_schedule_attributes
            )
        )
        await message.answer(current_schedule_by_attributes_existence_check_fail_chat_error_message)
        return

    schedule = current_schedule.schedule
    lessons: django.db.models.QuerySet = schedule.lessons.all()
    lessons.order_by('ordinal')

    date_format = NativeDateFormat.date_format
    today_string = today.strftime(date_format)

    poll_configuration = {'options': list(presence_option_to_string(o) for o in PresencePollOptions),
                           'type': 'quiz', 'correct_option_id': 0,
                           'is_anonymous': False,
                           'allows_multiple_answers': False,
                           'protect_content': True}

    if mode == PresenceMode.LIGHT_MODE:
        last_lesson = lessons[-1]
        try:
            poll__time_interval = poll_time_interval(mode, last_lesson=last_lesson)

        except LessonSkippedException: #TODO: add checker
            logger.info(lesson_time_skipped_logging_info_message.format(lesson_attributes="\bs"))
            await message.answer(lesson_time_skipped_chat_info_message.format(lesson_attributes=""))
            return

        deadline = poll__time_interval.upper

        question = f"{today_string} {chat_messages.presence_kw}"
        poll_identifying_attributes = {'question': question}
        poll_configuration.update(poll_identifying_attributes)

    if not mode == PresenceMode.LIGHT_MODE:
        await initiate_today_report(today, journal, mode)

        for lesson in lessons:
            lesson_name = f"{lesson.ordinal}-{lesson.subject.name}"

            await initiate_today_entries(today, journal, lesson)

            try:
                poll__time_interval = poll_time_interval(mode, lesson)

            except LessonSkippedException: #TODO: add checker
                lesson_name = f"{lesson.ordinal}-{lesson.subject.name}"
                await message.answer(lesson_time_skipped_chat_info_message.format(lesson_attributes=lesson_name))
                logger.info(
                    lesson_time_skipped_logging_info_message.format(lesson_attributes=vars(lesson))
                )
                continue

            poll_time = poll__time_interval.lower
            deadline = poll__time_interval.upper

            question = today_string + f" Заняття {lesson_name}"
            poll_identifying_attributes = {'question': question}
            poll_configuration.update(poll_identifying_attributes)

            till_poll = poll_time - timezone.localtime(timezone.now())
            till_poll_seconds = till_poll.seconds
            logger.info(
                presence_poll_on_send_expectation_logging_info_message.format(
                    poll_attributes = poll_identifying_attributes,
                    time=f"{till_poll_seconds} {seconds_kw}")

            )
            await asyncio.sleep(till_poll_seconds)
            till_deadline = deadline - timezone.localtime(timezone.now()) #TODO: create an async scheduler
            till_deadline_seconds = till_deadline.seconds
            poll_message = await message.answer_poll(**poll_configuration)
            poll_id = poll_message.poll.id
            poll_identifying_attributes.update({'id': poll_id})
            logger.info(
                presence_poll_on_sending_logging_info_message.format(
                    poll_attributes = poll_identifying_attributes
                )
            )
            await add_presence_poll(poll_id)
            logger.info(
                presence_poll_on_stop_expectation_logging_info_message.format(
                    poll_attributes = poll_identifying_attributes,
                    time=f'{till_deadline_seconds} {seconds_kw}')
            )
            await asyncio.sleep(till_deadline_seconds)  #TODO: schedule instead
            await bot.stop_poll(chat_id=poll_message.chat.id, message_id=poll_message.message_id)
            logger.info(
                presence_poll_on_stopping_logging_info_message.format(
                    poll_attributes = poll_identifying_attributes
                )
            )
            await delete_presence_poll(poll_id)
            logger.info(
                presence_poll_deleted_logging_info_message.format(poll_attributes=poll_id)
            )

        await amend_statuses(today, chat_id)
        logger.info(
            statuses_amended_logging_info_message
        )

    else:
        for lesson in lessons:
            await initiate_today_entries(today, journal, lesson)

        await initiate_today_report(today, journal, mode='L')

        logger.info(
            presence_poll_on_send_expectation_logging_info_message.format(
                poll_attributes=poll_identifying_attributes,
                time=0
            )
        )
        poll_message = await message.answer_poll(**poll_configuration)
        logger.info(
            presence_poll_on_sending_logging_info_message
        )
        poll_id = poll_message.poll.id
        poll_identifying_attributes.update({'id': poll_id})
        await add_presence_poll(poll_id)
        logger.info(
            presence_poll_added_logging_info_message.format(poll_id)
        )
        till_deadline = deadline - now
        till_deadline_seconds = till_deadline.seconds
        logger.info(
            presence_poll_on_stop_expectation_logging_info_message.format(time=till_deadline_seconds)
        )
        await asyncio.sleep(till_deadline_seconds) #TODO: schedule instead
        await bot.stop_poll(chat_id=poll_message.chat.id, message_id=poll_message.message_id)
        await delete_presence_poll(poll_id) #TODO: consider deleting presence polls
        logger.info(
            presence_poll_deleted_logging_info_message.format(poll_id)
        )
        logger.info(
            presence_poll_on_stopping_logging_info_message.format(
                poll_attrtibutes = poll_identifying_attributes
            )
        )

        await amend_statuses(today, chat_id)

    return


@commands_router.message(Command(commands='cancel', prefix=prefixes))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def cancel_command(message: types.Message, state: FSMContext, *args, **kwargs):
    current_state = await state.get_state()
    states_canceling_messages = {
        AbsenceReasonStates.AbsenceReason:
            [
                absence_reason_share_on_canceling_chat_info_message,
                absence_reason_share_on_canceling_logging_info_message
            ],
        SetJournalStatesGroup.set_journal:
            [
                journal_unset_success_chat_info_message,
                journal_unset_success_logging_info_message
            ],
        UserInformStatesGroup.receiver_id:
            [
                call_on_canceling_chat_info_message,
                call_on_canceling_logging_info_message
            ],
        GroupInformStatesGroup.receiver_id:
            [
                group_inform_on_canceling_chat_info_message,
                group_inform_on_canceling_logging_info_message
            ],
        NewScheduleStatesGroup.lessons_ordinals_to_subjects_names:
            [
                schedule_creation_on_cancel_chat_info_message,
                schedule_creation_on_cancel_logging_info_message
            ]
    }

    for state_key, (on_canceling_chat_message, on_canceling_logging_message) in states_canceling_messages.items():
        if current_state == state_key.state:
            response_on_canceling_chat_message = on_canceling_chat_message
            response_on_canceling_logging_message = on_canceling_logging_message
            break

    else:
        if current_state is None:
            response_on_canceling_chat_message = state_check_fail_chat_info_message
            response_on_canceling_logging_message = state_check_fail_logging_info_message

        else:
            logging.exception(
                data_input_on_canceling_logging_info_message
            )
            await message.reply(data_input_on_canceling_chat_info_message)
            return

    await state.clear()
    logger.info(response_on_canceling_logging_message)
    await message.reply(response_on_canceling_chat_message)


@commands_router.message(AbsenceReasonStates.AbsenceReason, F.text.regexp(r'Т'),
                         NoCommandFilter(),
                         F.chat.type.in_({'private'}))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def absence_reason_handler_T(message: types.Message, forms: FormsManager, *args, **kwargs): #TODO: rename to "suggestion handler" type
    await forms.show('absenceform')

@commands_router.message(AbsenceReasonStates.AbsenceReason, #TODO: rename to "suggestion handler" type
                         NoCommandFilter(),
                         F.text.regexp(r'Н'), F.chat.type.in_({'private'}))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def absence_reason_handler_H(message: types.Message, state: FSMContext, *args, **kwargs):
    await message.answer(absence_reason_share_on_canceling_chat_info_message)
    await state.clear()

@commands_router.message(AbsenceReasonStates.AbsenceReason, F.text.regexp(r'[^ТН]'), #TODO: rename to "suggestion handler" type
                         NoCommandFilter(),
                         F.chat.type.in_({'private'}))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def absence_reason_handler_invalid(message: types.Message, state: FSMContext, *args, **kwargs): #TODO: remove state argument
    await message.answer(absence_reason_share_suggestion_chat_field_message)
    logger.info(
        absence_reason_share_suggestion_logging_field_message
    )

@presence_poll_router.poll_answer(PresencePollFilter()) #TODO: add a flag for vote-answer mode, add an every-lesson mode
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def presence_poll_answer_handler (poll_answer: types.poll_answer, state: FSMContext, *args, **kwargs):  #TODO: add an ability to re-answer
    is_present = poll_answer.option_ids == [PresencePollOptions.Present.value]
    user_id = poll_answer.user.id

    logger.info(
        presence_poll_on_answer_logging_info_message.format(
            poll_attributes={"id": poll_answer.poll_id},
            answer_attributes=f"{poll_answer.option_ids[0]}:{is_present}"
        )
    )

    await process_user_on_lesson_presence(is_present, user_id)

    if not is_present:
        await bot.send_message(user_id, absence_reason_share_suggestion_chat_field_message)
        await state.set_state(AbsenceReasonStates.AbsenceReason)
        logger.info(
            absence_reason_share_suggestion_logging_field_message
        )


@commands_router.message(Command(commands=['absence_reason', 'ar'], prefix=prefixes), F.chat.type.in_({'private'}))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def absence_reason_command(message: types.Message, forms: FormsManager, *args, **kwargs):
    user_id = message.from_user.id
    #TODO: pass the lesson if lesson is none, then answer and return
    try:
        await validate_during_lesson_presence(user_id)

    except:
        await message.answer(during_lesson_check_fail_chat_error_message) #validation fail if of-lesson entry does not exist
        return

    if not await on_lesson_presence_check(user_id):
        await forms.show('absenceform')

    else:
        await message.answer(is_absent_check_fail_chat_error_message)


@registration_router.message(SuperuserKeyStates.key, F.chat.type.in_({'private'}))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def register_superuser_handler(message: types.Message, state: FSMContext, *args, **kwargs): #TODO: fix superuser multiple-registrations possibility
    user_id = message.from_user.id
    authentic_key = await state.get_data()

    try:
        validate_super_user_key(message.text, authentic_key['key'], user_id)

    except:
        await message.answer(key_validation_fail_chat_error_message)
        return

    try:
        await add_superuser(user_id)

    except Exception as e:
        logger.error(superuser_registration_fail_logging_error_message)
        await message.answer(registration_fail_chat_error_message)
        await state.clear()
        return

    logger.info(superuser_registration_success_logging_info_message)
    await message.answer(superuser_registration_success_chat_message)
    await state.clear()

async def request_key(chat_id: int, key: str) -> None:
    await bot.send_message(chat_id, key_field_chat_message)
    logger.info(key_field_logging_message.format(key=key))

@registration_router.message(Command(commands='register_superuser', prefix=prefixes),
                         F.chat.type.in_({'private'}),
                         RegisteredExternalIdFilter(Superuser))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def register_superuser_command(message: types.Message, state: FSMContext, *args, **kwargs):
    user_id = message.from_user.id

    key = key_generator.generate().get_key()

    await state.set_state(SuperuserKeyStates.key)
    await state.update_data(key=key)
    await request_key(user_id, key)


@registration_router.message(Command(commands='register', prefix=prefixes),
                             F.chat.type.in_({'private'}),
                             AftercommandFullCheck(allow_no_argument=True, modes=RegistrationMode, mode_checking=True),
                             RegisteredExternalIdFilter(Profile), SuperuserCalledUserToDELETEFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def register_command(message: types.Message, forms: FormsManager, state: FSMContext, mode=None, *args, **kwargs):
    user_id = message.from_user.id

    if mode == RegistrationMode.DELETE.value:
        try:
            if await is_superuser(user_id):
                called_user_id = await state.get_data()
                await delete_profile(called_user_id['Interlocutor_id'])
                logging.info(profile_deletion_success_chat_info_message)

            else:
                await delete_profile(user_id)
            await message.answer(text=profile_deletion_success_chat_info_message)


        except Exception:
            logging.error(profile_deletion_fail_logging_error_message, exc_info=True)
            await message.answer(text=profile_deletion_fail_chat_error_message)

    else:
        await forms.show('profileform')


register_journal_command_filters_config = (Command(commands='register_journal',prefix=prefixes),
                                           AftercommandFullCheck(allow_no_argument=True, modes=RegistrationMode, mode_checking=True),
                                           RegisteredExternalIdFilter(Journal, use_chat_id=True))


@registration_router.message(JournalRegistrationStates.key, NoCommandFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def register_journal_handler(message: types.Message, forms: FormsManager, state: FSMContext, *args, **kwargs):
    user_id = message.from_user.id
    chat_id = message.chat.id
    data = await state.get_data()
    authentic_key = data['key']
    try:
        validate_super_user_key(message.text, authentic_key, user_id)

    except:
        await message.answer(key_validation_fail_chat_error_message.format(user_id, chat_id))
        return

    set_journal = data['set_journal']
    journal = set_journal if set_journal else await get_journal_async({'external_id': chat_id})
    mode = data['mode']
    if mode == RegistrationMode.DELETE.value: #TODO: fix deletion, if no journal group journal (causes MultipleObjectsReturned)
        try:
            await delete_journal(journal)
            await message.answer(text=journal_deletion_success_chat_info_message)

        except Exception:
            await message.answer(text=journal_existence_check_fail_chat_message)

    else:
        await forms.show('journalform')

@journal_registration_subrouter.message(*register_journal_command_filters_config,
                         F.chat.type.in_({'private'}),
                         IsSuperUserFilter())
@journal_registration_subrouter.message(*register_journal_command_filters_config,
                         F.chat.type.in_({'group', 'supergroup'}),
                         IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def register_journal_command(
        message: types.Message    ,
        state: FSMContext         ,
        mode=None                 ,
        set_journal: Journal=None ,
        *args                     ,
        **kwargs                  ,
):
    chat_id = message.chat.id

    await state.set_state(JournalRegistrationStates.mode)
    await state.update_data(mode=mode)

    await state.set_state(JournalRegistrationStates.set_journal)
    await state.update_data(set_journal=set_journal)

    await state.set_state(JournalRegistrationStates.key)
    key = key_generator.generate().get_key()
    await state.update_data(key=key)
    await request_key(chat_id, key)
#TODO: catch report commands if no group journal
#TODO: make switches for report commands
today_report_command_filters_config = (Command(commands=['today_report', 'tr'], prefix=prefixes),
                                       AftercommandFullCheck(allow_no_argument=True, modes=ReportMode, flag_checking=True))

@journal_router.message(*today_report_command_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*today_report_command_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def today_report_command(
        message: types.Message    ,
        flag=ReportMode.Flag.TEXT ,
        set_journal: Journal=None ,
        *args                     ,
        **kwargs
):
    group_id = message.chat.id if not set_journal else set_journal.external_id
    try:
        today_report = await get_on_mode_report(group_id, ReportMode.TODAY)

    except Exception:
        await message.answer(report_parameters_check_fail_chat_error_message)
        return

    table = await make_report_table(today_report)
    summary = await report_summary(today_report, ReportMode.TODAY)

    date_format = NativeDateFormat.date_format
    date_string = today_report.date.strftime(date_format)

    await message.answer(
        report_description_chat_info_message.format(report_parameters=f"{chat_messages.on_kw} {date_string}")
    )

    match ReportMode.Flag(flag):

        case ReportMode.Flag.TEXT:
            await message.answer(f"```{table}```", 'Markdown')
            await message.answer(f"```{summary}```", 'Markdown', disable_notification=True)

        case ReportMode.Flag.DOCUMENT:
                temp_path = os.path.join(tempfile.gettempdir(), os.urandom(24).hex()) + '.docx'
                document = docx.Document()
                parser = htmldocx.HtmlToDocx()

                table_html = table.get_html_string()
                parser.add_html_to_document(table_html, document)
                document.save(temp_path)
                input_file = types.FSInputFile(temp_path)
                await message.answer_document(input_file)

                document._body.clear_content()


                summary_html = summary.get_html_string()
                parser.add_html_to_document(summary_html, document)
                document.save(temp_path)
                input_file = types.FSInputFile(temp_path)
                await message.answer_document(input_file, disable_notification=True)


last_report_command_filters_config = (Command(commands=['last_report', 'lr'], prefix=prefixes),
                                      AftercommandFullCheck(allow_no_argument=True, modes=ReportMode, flag_checking=True))

@journal_router.message(*last_report_command_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*last_report_command_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def last_report_command(
        message: types.Message    ,
        flag=ReportMode.Flag.TEXT ,
        set_journal: Journal=None ,
        *args                     ,
        **kwargs
):
    group_id = message.chat.id if not set_journal else set_journal.external_id

    try:
        last_report = await get_on_mode_report(group_id, ReportMode.LAST)

    except Exception:
        await message.answer(report_parameters_check_fail_chat_error_message)
        return

    table = await make_report_table(last_report)
    summary = await report_summary(last_report, ReportMode.LAST)

    date_format = NativeDateFormat.date_format
    date_string = last_report.date.strftime(date_format)

    await message.answer(
        report_description_chat_info_message.format(report_parameters=f"{chat_messages.on_kw} {date_string}")
    )

    match ReportMode.Flag(flag):

        case ReportMode.Flag.TEXT:
            await message.answer(f"```{table}```", 'Markdown')
            await message.answer(f"```{summary}```", 'Markdown', disable_notification=True)

        case ReportMode.Flag.DOCUMENT:
            temp_path = os.path.join(tempfile.gettempdir(), os.urandom(24).hex()) + '.docx'
            document = docx.Document()
            parser = htmldocx.HtmlToDocx()

            table_html = table.get_html_string()
            parser.add_html_to_document(table_html, document)
            document.save(temp_path)
            input_file = types.FSInputFile(temp_path)
            await message.answer_document(input_file)

            document._body.clear_content()

            summary_html = summary.get_html_string()
            parser.add_html_to_document(summary_html, document)
            document.save(temp_path)
            input_file = types.FSInputFile(temp_path)
            await message.answer_document(input_file, disable_notification=True)


on_date_report_command_filters_config = (Command(commands=['on_date_report', 'odr'], prefix=prefixes),
                                         AftercommandFullCheck(allow_no_argument=False,
                                                        modes=ReportMode,
                                                        additional_arguments_checker=date_validator,
                                                        flag_checking=True))
@journal_router.message(*on_date_report_command_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*on_date_report_command_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def on_date_report_command(
        message: types.Message     ,
        additional_arguments=False ,
        flag=ReportMode.Flag.TEXT  ,
        set_journal: Journal=None  ,
        *args                      ,
        **kwargs
):
    date_string = additional_arguments[0] #TODO: fix additional_arguments duplication
    date_format = NativeDateFormat.date_format
    date = datetime.datetime.strptime(date_string, date_format).date()

    group_id = message.chat.id if not set_journal else set_journal.external_id

    try:
        on_date_report = await get_on_mode_report(group_id, ReportMode.ON_DATE, date)

    except: #TODO: write a decorator-validator instead
        await message.answer(report_parameters_check_fail_chat_error_message)
        return

    table = await make_report_table(on_date_report)
    summary = await report_summary(on_date_report, ReportMode.ON_DATE)

    await message.answer(
        report_description_chat_info_message.format(report_parameters=f"{chat_messages.on_kw} {date_string}")
    )

    match ReportMode.Flag(flag):

        case ReportMode.Flag.TEXT:
            await message.answer(f"```{table}```", 'Markdown')
            await message.answer(f"```{summary}```", 'Markdown', disable_notification=True)

        case ReportMode.Flag.DOCUMENT:
            temp_path = os.path.join(tempfile.gettempdir(), os.urandom(24).hex()) + '.docx'
            document = docx.Document()
            parser = htmldocx.HtmlToDocx()

            table_html = table.get_html_string()
            parser.add_html_to_document(table_html, document)
            document.save(temp_path)
            input_file = types.FSInputFile(temp_path)
            await message.answer_document(input_file)

            document._body.clear_content()

            summary_html = summary.get_html_string()
            parser.add_html_to_document(summary_html, document)
            document.save(temp_path)
            input_file = types.FSInputFile(temp_path)
            await message.answer_document(input_file, disable_notification=True)


all_reports__command_filters_config = (Command(commands=('allreport', 'all_reports'), prefix=prefixes), )
@journal_router.message(*all_reports__command_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*all_reports__command_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def allreport_command(message: Message, set_journal: Journal=None, *args, **kwargs): #TODO: consider making an analogue uni-table
    group_id = message.chat.id
    journal = set_journal if set_journal else await get_journal_async({'external_id': group_id})
    all_journal_report_parameters = ReportParameters.objects.filter(journal=journal)

    allreport_temporary_filepath = os.path.join(tempfile.gettempdir(), os.urandom(24).hex()) + '.docx'
    allreport_document = docx.Document()
    html_to_docx__parser = htmldocx.HtmlToDocx()

    for report_parameters in all_journal_report_parameters:
        report_table = await make_report_table(report_parameters)
        if report_table:
            report_table_html = report_table.get_html_string()
            html_to_docx__parser.add_html_to_document(f"<center>{report_parameters.date}</center>", allreport_document)
            html_to_docx__parser.add_html_to_document(report_table_html, allreport_document)
            html_to_docx__parser.add_html_to_document("<span></span><span></span>   ", allreport_document)

    allreport_document.save(allreport_temporary_filepath)

    input_file = types.FSInputFile(allreport_temporary_filepath)
    await message.answer_document(input_file)

    allreport_document._body.clear_content()


dossier_command_filters_config = (Command(commands='dossier', prefix=prefixes),
                                  AftercommandFullCheck(allow_no_argument=True,
                                                         modes=ReportMode,
                                                         allow_no_mode=True,
                                                         flag_checking=True))

@journal_router.message(*dossier_command_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*dossier_command_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def dossier_command(message: Message, flag=ReportMode.Flag.TEXT,set_journal: Journal=None, *args, **kwargs):
    #TODO: add dossier for profile, function will work depending on chat type
    group_id = message.chat.id if not set_journal else set_journal.external_id

    table = await get_journal_dossier(group_id)

    match ReportMode.Flag(flag):
        case ReportMode.Flag.TEXT:
            await message.answer(f"```{table}```", 'Markdown')

        case ReportMode.Flag.DOCUMENT:
            temp_path = os.path.join(tempfile.gettempdir(), os.urandom(24).hex()) + '.docx'
            document = docx.Document()
            parser = htmldocx.HtmlToDocx()

            table_html = table.get_html_string()
            parser.add_html_to_document(table_html, document)
            document.save(temp_path)
            input_file = types.FSInputFile(temp_path)
            await message.answer_document(input_file)


today_schedule_command_filters_config = (
    Command(commands=['today_schedule', 'ts'], prefix=prefixes) ,
    AftercommandFullCheck(
        allow_no_argument=True,
        modes=ReportMode,
        allow_no_mode=True,
        flag_checking=True,
    )
)
@journal_router.message(*today_schedule_command_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*today_schedule_command_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def today_schedule__command(
        message: Message                         ,
        flag: enum.auto|str=ReportMode.Flag.TEXT , #TODO: check annotation accuracy
        set_journal: Journal=None                ,
        *args, **kwargs                          ,
) -> None:
    group_id = message.chat.id if not set_journal else set_journal.external_id
    journal = set_journal if set_journal else await get_journal_async({'external_id': group_id})

    today_current_schedule = await get_on_mode_journal_current_schedule(journal, ReportMode.TODAY)
    today_current_schedule_schedule = today_current_schedule.schedule
    today_current_schedule_schedule_table = make_schedule_table(today_current_schedule_schedule)

    match ReportMode.Flag(flag):
        case ReportMode.Flag.TEXT:
            await message.answer(f"```{today_current_schedule_schedule_table}```", 'Markdown')

        case ReportMode.Flag.DOCUMENT:
            temporary_path = os.path.join(tempfile.gettempdir(), os.urandom(24).hex()) + '.docx'
            document = docx.Document()
            parser = htmldocx.HtmlToDocx()

            table_html = today_current_schedule_schedule_table.get_html_string()
            parser.add_html_to_document(table_html, document)
            document.save(temporary_path)
            input_file = types.FSInputFile(temporary_path)
            await message.answer_document(input_file)


last_schedule_command_filters_config = (
    Command(commands=['last_schedule', 'ls'], prefix=prefixes) ,
    AftercommandFullCheck(
        allow_no_argument=True,
        modes=ReportMode,
        allow_no_mode=True,
        flag_checking=True
    )
)
@journal_router.message(*last_schedule_command_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*last_schedule_command_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def last_schedule__command(
        message: Message                         ,
        flag: enum.auto|str=ReportMode.Flag.TEXT , #TODO: check annotation accuracy
        set_journal: Journal=None                ,
        *args, **kwargs                          ,
) -> None:
    group_id = message.chat.id if not set_journal else set_journal.external_id
    journal = set_journal if set_journal else await get_journal_async({'external_id': group_id})

    last_current_schedule = await get_on_mode_journal_current_schedule(journal, ReportMode.LAST)
    last_current_schedule_schedule = last_current_schedule.schedule
    last_current_schedule_schedule_table = make_schedule_table(last_current_schedule_schedule)

    await message.answer(
        schedule_description_chat_info_message.format(
            schedule_parameters=f"{chat_messages.on_kw} {last_current_schedule.date}"
        )
    )

    match ReportMode.Flag(flag):
        case ReportMode.Flag.TEXT:
            await message.answer(f"```{last_current_schedule_schedule_table}```", 'Markdown', disable_notification=True)

        case ReportMode.Flag.DOCUMENT:
            temporary_path = os.path.join(tempfile.gettempdir(), os.urandom(24).hex()) + '.docx'
            document = docx.Document()
            parser = htmldocx.HtmlToDocx()

            table_html = last_current_schedule_schedule_table.get_html_string()
            parser.add_html_to_document(table_html, document)
            document.save(temporary_path)
            input_file = types.FSInputFile(temporary_path)
            await message.answer_document(input_file, disable_notification=True)


on_date_schedule_command_filters_config = (
    Command(commands=['on_date_schedule', 'ods'], prefix=prefixes) ,
    AftercommandFullCheck(
        allow_no_argument=False,
        modes=ReportMode,
        allow_no_mode=True,
        flag_checking=True,
        additional_arguments_checker=date_validator
    )
)
@journal_router.message(*on_date_schedule_command_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*on_date_schedule_command_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def on_date_schedule__command(
        message: Message                         ,
        flag: enum.auto|str=ReportMode.Flag.TEXT , #TODO: check annotation accuracy
        additional_arguments: list[str]=None     ,
        set_journal: Journal=None                ,
        *args, **kwargs                          ,
) -> None:
    specified_date_string = additional_arguments[0]  # TODO: fix additional_arguments duplication
    native_date_format = NativeDateFormat.date_format
    date = datetime.datetime.strptime(specified_date_string, native_date_format).date()
    group_id = message.chat.id if not set_journal else set_journal.external_id
    journal = set_journal if set_journal else await get_journal_async({'external_id': group_id})

    last_current_schedule = await get_on_mode_journal_current_schedule(journal, ReportMode.ON_DATE, date)
    last_current_schedule_schedule = last_current_schedule.schedule
    last_current_schedule_schedule_table = make_schedule_table(last_current_schedule_schedule)

    match ReportMode.Flag(flag):
        case ReportMode.Flag.TEXT:
            await message.answer(f"```{last_current_schedule_schedule_table}```", 'Markdown')

        case ReportMode.Flag.DOCUMENT:
            temporary_path = os.path.join(tempfile.gettempdir(), os.urandom(24).hex()) + '.docx'
            document = docx.Document()
            parser = htmldocx.HtmlToDocx()

            table_html = last_current_schedule_schedule_table.get_html_string()
            parser.add_html_to_document(table_html, document)
            document.save(temporary_path)
            input_file = types.FSInputFile(temporary_path)
            await message.answer_document(input_file)


@journal_router.message(SubjectRegistrationStatesGroup.subject_registration, NoCommandFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def subject_registration_handler(message: Message, state: FSMContext, set_journal: Journal=None):
    subject_name = message.text

    if not await check_subject_is_not_created(subject_name):
        await message.answer(subject_is_not_created_by_attributes_check_fail_chat_error_message.format(
            subject_attributes=chat_messages.name_obj_kw
            )
        )
        return

    try:
        subject_attributes = {'name': subject_name, 'journal': set_journal}
        await add_subject(subject_attributes)

    except Exception: #TODO: define exception maintaining lack reason
        await message.answer(subject_creation_fail_logging_error_message)
        logging.error(subject_creation_fail_logging_error_message)
        await state.clear()
        return

    logger.info(
        subject_creation_success_logging_info_message.format(subject_attributes=subject_attributes)
    )
    await message.answer(subject_creation_success_chat_info_message)
    await state.clear()


add_subject_filters_config = {Command(commands=['create_subject'])}

@journal_router.message(*add_subject_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*add_subject_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def create_subject_command(message: Message, state: FSMContext, set_journal: Journal=None, *args, **kwargs):

    await message.answer(subject_name_chat_field_message)
    logging.info(subject_name_logging_field_message)

    await state.set_state(SubjectRegistrationStatesGroup.set_journal)
    if set_journal:
        await state.update_data(set_journal=set_journal)

    else:
        group_id = message.chat.id
        journal_attributes = {'external_id': group_id}
        journal = await get_journal_async(journal_attributes)
        await state.update_data(set_journal=journal)

    await state.set_state(SubjectRegistrationStatesGroup.subject_registration)


@journal_router.message(PostScheduleStatesGroup.schedule_id, NoCommandFilter(), ScheduleExistsFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def post_schedule_handler(message: Message, state: FSMContext):
    try:
        schedule = await get_schedule_async(id=message.text)
        state_data = await state.get_data()
        try:
            set_journal = state_data['set_journal']
            journal = set_journal

        except KeyError:
            journal = await get_journal_async({'external_id': message.chat.id})

        date = state_data['date']
        current_schedule_attributes = {"schedule": schedule, "journal": journal, "date": date}
        await add_current_schedule(current_schedule_attributes)

    except Exception:
        logger.error(
            current_schedule_post_fail_logging_error_message.format(
                current_schedule_attributes=current_schedule_attributes
            )
        )
        await message.answer(current_schedule_post_fail_chat_error_message)

        await state.clear()

        return

    logger.info(
        current_schedule_post_success_logging_info_message.format(
            current_schedule_attributes=current_schedule_attributes
        )
    )
    await message.answer(current_schedule_post_success_chat_info_message)

    await state.clear()

post_schedule_filters_config = (Command(commands=['post_schedule', 'ps']),
                                AftercommandFullCheck(allow_no_argument=False,
                                                      additional_arguments_checker=date_validator))

@journal_router.message(*post_schedule_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*post_schedule_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def post_schedule_command(
        message: Message           ,
        state: FSMContext          ,
        additional_arguments=False ,
        set_journal: Journal=None  ,
        *args                      ,
        **kwargs
):
    date_string = additional_arguments[0]
    date_format = NativeDateFormat.date_format
    date = datetime.datetime.strptime(date_string, date_format).date()

    await state.set_state(PostScheduleStatesGroup.date)
    await state.update_data(date=date)

    if set_journal:
        journal = set_journal
        await state.set_state(PostScheduleStatesGroup.set_journal)
        await state.update_data(set_journal=journal)

    await state.set_state(PostScheduleStatesGroup.schedule_id)
    await message.answer(schedule_id_chat_field_message)

@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def schedule_completion_handler(message: Message, state: FSMContext, *args, **kwargs): #TODO: add schedule existence checking
    state_data = await state.get_data()
    try:
        lessons_ordinals_to_subjects_names: dict = state_data["lessons_ordinals_to_subjects_names"]

    except KeyError:
        lessons_ordinals_to_subjects_names = {}

    if not check_schedule_dict_is_not_empty(lessons_ordinals_to_subjects_names):
        await message.answer(schedule_is_not_empty_check_fail_chat_error_message)

    else:
        try:
            schedule = await add_schedule(lessons_ordinals_to_subjects_names)

        except Exception as e:
            raise e

        await message.answer(
            schedule_creation_success_chat_info_message.format(
                schedule_attributes=schedule.id
            )
        )

    await state.clear()


@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
@commands_router.message(NewScheduleStatesGroup.lesson_ordinal, NoCommandFilter())
async def schedule_lesson_setting_suggest(message: Message, state: FSMContext, *args, **kwargs):
    lesson_ordinal = (await state.get_data())['lesson_ordinal']
    await message.answer(lesson_select_chat_field_message.format(lesson_attributes=lesson_ordinal))
    logger.info(lesson_select_logging_field_message.format(lesson_attributes=lesson_ordinal))
    await state.set_state(NewScheduleStatesGroup.lessons_ordinals_to_subjects_names)

schedule_lesson_setting_handler_filters_config = (NewScheduleStatesGroup.lessons_ordinals_to_subjects_names, NoCommandFilter())

@commands_router.message(*schedule_lesson_setting_handler_filters_config, SubjectExistsFilter())
@commands_router.message(*schedule_lesson_setting_handler_filters_config, F.text.regexp(f'{void_symbol}'))
@commands_router.message(*schedule_lesson_setting_handler_filters_config, F.text.regexp(f'{finish_word}'))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def schedule_lesson_setting_handler(message: Message, state: FSMContext, *args, **kwargs):
    state_data = await state.get_data()
    lesson_ordinal = state_data['lesson_ordinal']

    scheduling_keyword = message.text

    if scheduling_keyword == finish_word:
        if lesson_ordinal == 1: #TODO: consider block importance
            logger.error(schedule_is_not_empty_check_fail_logging_error_message)
            await message.answer(schedule_is_not_empty_check_fail_chat_error_message)
            await state.clear()
            return

        await schedule_completion_handler(message, state)
        return

    if scheduling_keyword != void_symbol:
        lessons_ordinals_to_subjects_names = state_data['lessons_ordinals_to_subjects_names']
        complete_lessons_ordinals_to_subjects_names = lessons_ordinals_to_subjects_names
        complete_lessons_ordinals_to_subjects_names[lesson_ordinal] = scheduling_keyword
        await state.update_data(lessons_ordinals_to_subjects_names=complete_lessons_ordinals_to_subjects_names)

    if lesson_ordinal < ScheduleTiming.last_lesson_ordinal:
        await state.set_state(NewScheduleStatesGroup.lesson_ordinal)
        await state.update_data(lesson_ordinal=lesson_ordinal + 1)
        await schedule_lesson_setting_suggest(message, state)

    else:
        await state.set_state(NewScheduleStatesGroup.complete)
        await schedule_completion_handler(message, state)

create_schedule_command_filters_config = (Command(commands=['create_schedule', 'ns'], prefix=prefixes), DummyFilter()) #TODO: consider dummy-filter

@commands_router.message(*create_schedule_command_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@commands_router.message(*create_schedule_command_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def create_schedule_command(message: Message, state: FSMContext, *args, **kwargs):
    await message.answer(schedule_building_instruction)
    await state.update_data(lessons_ordinals_to_subjects_names={})
    await state.set_state(NewScheduleStatesGroup.lesson_ordinal)
    await state.update_data(lesson_ordinal=1)
    await asyncio.sleep(messaging_pause)
    await schedule_lesson_setting_suggest(message, state)


@commands_router.message(Command(commands=['schedule', 's'], prefix=prefixes))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def schedule_command(message: Message, state: FSMContext, *args, **kwargs):
    raise NotImplementedError


@journal_router.message(ReportRedoStatesGroup.redoing, NoCommandFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def redo_report_handler(message: Message, state: FSMContext, set_journal: Journal=None, *args, **kwargs):
    group_id = message.chat.id if not set_journal else set_journal.external_id
    journal = await get_journal_async({'external_id': group_id}) if not set_journal else set_journal

    try:
        data = await state.get_data()
        date = data['date']

        slash_n_free_message_text = message.text.replace("\n", "")

        try:
            validate_report_format(slash_n_free_message_text)

        except Exception:
            await message.answer(report_table_format_validation_chat_error_message)
            return

        if await validate_report_name_references(slash_n_free_message_text, journal): #TODO: change if to try
            report = slash_n_free_message_text
            report_headers_match = re.search(report_headers_rePattern, report)
            lessons_ordinals_string = report_headers_match.group(2)
            lessons_ordinals_string_split: list = lessons_ordinals_string.split()
            lessons_ordinals: list = [int(l) for l in lessons_ordinals_string_split]
            report_rows = regex.finditer(report_row_rePattern, report)
            for row_match in report_rows:
                await redo_entries_by_report_row(row_match, group_id, date, lessons_ordinals)

            logger.info(report_redo_success_logging_info_message)
            await message.answer(report_redo_success_chat_info_message)

        else:
            await message.answer(report_table_name_references_validation_fail_chat_error_message)
            return

    except Exception as e:
        await message.answer(report_redo_fail_chat_error_message)

redo_report_filters_config = (Command(commands=['redo_report', 'rr'], prefix=prefixes),
                              AftercommandFullCheck(
                                  allow_no_argument=False,
                                  allow_no_mode=False, #TODO: check-fix
                                  additional_arguments_checker=date_validator))  #TODO: add flags (show_current, get example)

@journal_router.message(*redo_report_filters_config, F.chat.type.in_({'private'}), IsSuperUserFilter())
@journal_router.message(*redo_report_filters_config, F.chat.type.in_({'group', 'supergroup'}), IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def redo_report_command(
        message: Message          ,
        state: FSMContext         ,
        additional_arguments      ,
        set_journal: Journal=None ,
        *args                     ,
        **kwargs
): #TODO: fix superusage state bug
    await message.answer(redo_report_chat_field_message)
    logger.info(redo_report_logging_field_message)
    await message.answer(f"```{report_example_text}```", parse_mode='Markdown')
    await state.set_state(ReportRedoStatesGroup.date)
    date_string = additional_arguments[0]
    date_format = NativeDateFormat.date_format
    date = datetime.datetime.strptime(date_string, date_format).date()
    await state.update_data(date=date)
    await state.set_state(ReportRedoStatesGroup.redoing)


@commands_router.message(SetJournalStatesGroup.setting_journal, NoCommandFilter(), F.chat.type.in_({'private'}))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def set_journal_handler(message: types.Message, state: FSMContext, *args, **kwargs):
    response = message.text
    try:
        validate_journal_format(response)

    except ValidationError as e:
        error_message = e.message
        await message.answer(error_message)
        return

    try:
        await check_journal_is_registered_by_name(response)

    except ValidationError as e:
        error_message = e.message
        await message.answer(error_message)
        return

    set_journal_name = response
    set_journal = await get_journal_async({'name': set_journal_name})

    await state.set_state(SetJournalStatesGroup.set_journal)
    await state.update_data(set_journal=set_journal)
    logger.info(journal_setting_success_logging_info_message.format(journal_attributes=set_journal.__dict__)) #TODO: fix model state attribute pass
    await message.answer(journal_set_success_chat_info_message)


async def request_journal(chat_id: int):
    await bot.send_message(chat_id, journal_name_chat_field_message)
    logger.info(journal_name_logging_field_message)

@commands_router.message(Command(commands=['set_journal', 'sj'], prefix=prefixes),
                         F.chat.type.in_({'private'}),
                         IsSuperUserFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False) #TODO: add journal name at journal-set-logging-message
async def set_journal_command(message: types.Message, state: FSMContext, *args, **kwargs):
    await request_journal(message.chat.id)
    await state.set_state(SetJournalStatesGroup.setting_journal)


@commands_router.message(UserInformStatesGroup.receiver_id, NoCommandFilter(), F.chat.type.in_({'private'}))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def user_inform_handler(message: types.Message, state: FSMContext, *args, **kwargs):
    user_id = message.from_user.id
    interlocutor_id = await state.get_data()
    await bot.send_message(interlocutor_id['Interlocutor_id'], message.text)
    await state.update_data(receiver_id=user_id)

@commands_router.message(UserInformStatesGroup.call, NoCommandFilter(), F.chat.type.in_({'private'}))
@log_track_frame(untracked_log_data, track_non_keyword_args=False)
async def user_call_handler(message: types.Message, state: FSMContext, *args, **kwargs):
    response = message.text

    try:
        validate_name_format(response)
    except Exception:
        await message.answer(profile_name_format_validation_fail_logging_error_message)
        return
    name = response
    try:
        profile = await get_profile_async(name=name)

    except Exception:
        await message.answer(
            profile_is_registered_by_attributes_сheck_fail_chat_error_message.format(
                profile_attributes=chat_messages.profile_name_kwс
            )
        )
        await state.clear()
        return

    profile_id = profile.external_id
    await state.set_state(UserInformStatesGroup.receiver_id)
    await state.update_data(Interlocutor_id=profile_id)
    await message.answer(inform_to_profile_user_chat_info_message.format(name))
    logging.info(inform_to_sender_logging_info_message)
    await bot.send_message(profile_id, user_inform_to_receiver_chat_error_message)

@commands_router.message(Command(commands=['call'], prefix=prefixes),
                         F.chat.type.in_({'private'}),
                         IsSuperUserFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def call_command(message: types.Message, state: FSMContext, *args, **kwargs):
    await state.set_state(UserInformStatesGroup.call)
    await message.answer(profile_name_logging_field_message)

async def inform_profile_user(profile: Profile, message_text=str) -> None:
    await bot.send_message(profile.external_id, message_text)

async def inform_journal_group(journal: Journal, message_text=str) -> None:
    await bot.send_message(journal.external_id, message_text)


async def inform_all_journal_profiles_users_and_journal_group(journal, message_text):
    all_journal_profiles = await get_all_journal_profiles(journal)
    for profile in all_journal_profiles:
        await inform_profile_user(profile, message_text)

    await inform_journal_group(journal, message_text)


@commands_router.message(GroupInformStatesGroup.receiver_id, NoCommandFilter(), F.chat.type.in_({'private'}))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def group_inform_handler(message: types.Message, state: FSMContext, *args, **kwargs):
    journal_external_id = await state.get_data()
    journal = await get_journal_async({'external_id': journal_external_id['receiver_id']})
    await inform_all_journal_profiles_users_and_journal_group(journal, message.text)
    await state.clear()

@commands_router.message(GroupInformStatesGroup.call, NoCommandFilter(), F.chat.type.in_({'private'}))
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def group_call_handler(message: types.Message, state: FSMContext, *args, **kwargs):
    response = message.text

    try:
        validate_journal_format(response)

    except Exception:
        await message.answer(journal_name_format_validation_fail_chat_error_message)
        return

    journal_name = response
    journal_attributes = {'name': journal_name}
    try:
        await validate_journal_is_registered(**journal_attributes)
        journal = await get_journal_async(journal_attributes) #TODO: add check

    except Exception:
        await message.answer(
            journal_is_registered_by_attributes_check_fail_chat_error_message.format(
                journal_attributes=chat_messages.journal_name_kw
            )
        )
        await state.clear()
        return

    await state.set_state(GroupInformStatesGroup.receiver_id)
    await state.update_data(receiver_id=journal.external_id)
    await message.answer(group_inform_chat_info_message.format(journal_name))
    await inform_all_journal_profiles_users_and_journal_group(journal, group_inform_to_receiver_chat_info_message)


@commands_router.message(Command(commands=['groupcall', 'gc'], prefix=prefixes),
                         F.chat.type.in_({'private'}),
                         IsSuperUserFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def group_call_command(message: types.Message, state: FSMContext, *args, **kwargs):
    await state.set_state(GroupInformStatesGroup.call)
    await message.answer(journal_name_chat_field_message)
    logger.info(journal_name_logging_field_message)


@commands_router.message(Command(commands=['leave_chat_delete_journal']),
                        F.chat.type.in_({'group', 'supergroup'}),
                        IsAdminFilter())
@log_track_frame(untracked_data=untracked_log_data, track_non_keyword_args=False)
async def leave_chat_delete_journal_command(message: types.Message, *args, **kwargs):
    group_id = message.chat.id

    await delete_journal(external_id=group_id)
    await message.answer(journal_deletion_success_chat_info_message)
    await bot.leave_chat(group_id)
    logger.info(bot_leaved_logging_info_message)


#TODO: add an execute command for admin


#TODO: add separate aftercommand-elements-entering-possibility mechanism